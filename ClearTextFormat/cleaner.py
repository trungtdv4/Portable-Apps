#!/usr/bin/env python3
"""
Text Cleaner — Khử nhiễu Unicode cho văn bản tiếng Việt (và mọi ngôn ngữ)
Hỗ trợ: .xlsx, .csv, .txt, .md, .docx
"""

import sys
import os
import re
import unicodedata
import shutil
from pathlib import Path

# ── Màu terminal (tắt tự động nếu không hỗ trợ) ──────────────────────────────
def supports_color():
    return hasattr(sys.stdout, "isatty") and sys.stdout.isatty()

C = {
    "reset":  "\033[0m"  if supports_color() else "",
    "bold":   "\033[1m"  if supports_color() else "",
    "green":  "\033[92m" if supports_color() else "",
    "yellow": "\033[93m" if supports_color() else "",
    "cyan":   "\033[96m" if supports_color() else "",
    "red":    "\033[91m" if supports_color() else "",
    "gray":   "\033[90m" if supports_color() else "",
}

def c(color, text):
    return f"{C[color]}{text}{C['reset']}"

# ── Core: logic khử nhiễu (port từ VBA v4) ───────────────────────────────────

# Bảng chuyển đổi Unicode Tổ Hợp → Dựng Sẵn cho tiếng Việt
# Mỗi entry: (base_char, combining_codepoint) → precomposed_char
_COMBINING_MAP = {}

def _build_combining_map():
    """Xây dựng bảng tổ hợp → dựng sẵn cho tất cả nguyên âm tiếng Việt."""
    # Base chars: a e i o u y + â ê ô ă ơ ư (thường và hoa)
    bases = list("aeiouyAEIOUY") + [
        "\u00e2","\u00ea","\u00f4","\u0103","\u01a1","\u01b0",  # â ê ô ă ơ ư
        "\u00c2","\u00ca","\u00d4","\u0102","\u01a0","\u01af",  # Â Ê Ô Ă Ơ Ư
    ]
    # 5 combining diacritics: huyền sắc hỏi ngã nặng
    combinings = ["\u0300","\u0301","\u0309","\u0303","\u0323"]

    for base in bases:
        for comb in combinings:
            composed = unicodedata.normalize("NFC", base + comb)
            if composed != base + comb:  # NFC thực sự ghép được
                _COMBINING_MAP[base + comb] = composed

_build_combining_map()

# Sắp xếp dài trước để tránh partial match
_COMBINING_PAIRS = sorted(_COMBINING_MAP.items(), key=lambda x: -len(x[0]))

# Regex blacklist: control chars + C1 block + combining marks còn sót
_JUNK_RE = re.compile(
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f"  # control chars (giữ \t \n \r)
    r"\x80-\x9f"                           # C1 control block
    r"\u0300-\u036f"                       # combining marks còn sót
    r"\u200b-\u200d"                       # zero-width chars
    r"\ufeff]"                             # BOM
)

_MULTI_SPACE_RE = re.compile(r"[ \t]{2,}")


def clean_text(text: str) -> str:
    """Khử nhiễu một chuỗi văn bản."""
    if not text:
        return text

    # 1. Đổi NBSP và tab → khoảng trắng thường
    text = text.replace("\u00a0", " ").replace("\t", " ")

    # 2. Convert tổ hợp → dựng sẵn (PHẢI trước regex)
    for combo, precomposed in _COMBINING_PAIRS:
        if combo in text:
            text = text.replace(combo, precomposed)

    # 3. Xóa ký tự rác (blacklist — giữ lại mọi thứ hợp lệ kể cả CJK, ký hiệu)
    text = _JUNK_RE.sub("", text)

    # 4. Chuẩn hóa khoảng trắng thừa (giữ newline)
    text = _MULTI_SPACE_RE.sub(" ", text)

    return text.strip()


def count_dirty_chars(text: str) -> int:
    """Đếm số ký tự nhiễu trong một chuỗi (để báo cáo)."""
    if not text:
        return 0
    count = 0
    for combo in _COMBINING_MAP:
        count += text.count(combo)
    count += len(_JUNK_RE.findall(text))
    return count


# ── Xử lý từng loại file ─────────────────────────────────────────────────────

def process_xlsx(src: Path, dst: Path) -> dict:
    try:
        import openpyxl
    except ImportError:
        raise ImportError("Thiếu thư viện: chạy  pip install openpyxl")

    wb = openpyxl.load_workbook(src)
    stats = {"cells": 0, "dirty": 0}

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    stats["cells"] += 1
                    dirty = count_dirty_chars(cell.value)
                    if dirty:
                        stats["dirty"] += 1
                    cell.value = clean_text(cell.value)

    wb.save(dst)
    return stats


def process_csv(src: Path, dst: Path) -> dict:
    import csv
    stats = {"cells": 0, "dirty": 0}

    # Thử detect encoding
    encodings = ["utf-8-sig", "utf-8", "cp1258", "latin-1"]
    content = None
    used_enc = "utf-8"
    for enc in encodings:
        try:
            content = src.read_text(encoding=enc)
            used_enc = enc
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        raise ValueError("Không đọc được file CSV — encoding không xác định")

    import io
    reader = csv.reader(io.StringIO(content))
    rows = []
    for row in reader:
        new_row = []
        for cell in row:
            stats["cells"] += 1
            dirty = count_dirty_chars(cell)
            if dirty:
                stats["dirty"] += 1
            new_row.append(clean_text(cell))
        rows.append(new_row)

    with open(dst, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.writer(f)
        writer.writerows(rows)

    return stats


def process_text(src: Path, dst: Path) -> dict:
    encodings = ["utf-8-sig", "utf-8", "cp1258", "latin-1"]
    content = None
    for enc in encodings:
        try:
            content = src.read_text(encoding=enc)
            break
        except UnicodeDecodeError:
            continue

    if content is None:
        raise ValueError("Không đọc được file — encoding không xác định")

    stats = {"cells": 0, "dirty": 0}
    lines = content.splitlines(keepends=True)
    cleaned_lines = []
    for line in lines:
        stats["cells"] += 1
        dirty = count_dirty_chars(line)
        if dirty:
            stats["dirty"] += 1
        # Giữ newline, chỉ clean phần text
        nl = ""
        stripped = line
        if line.endswith("\r\n"):
            nl, stripped = "\r\n", line[:-2]
        elif line.endswith("\n"):
            nl, stripped = "\n", line[:-1]
        elif line.endswith("\r"):
            nl, stripped = "\r", line[:-1]
        cleaned_lines.append(clean_text(stripped) + nl)

    dst.write_text("".join(cleaned_lines), encoding="utf-8")
    return stats


def process_docx(src: Path, dst: Path) -> dict:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("Thiếu thư viện: chạy  pip install python-docx")

    doc = Document(src)
    stats = {"cells": 0, "dirty": 0}

    def clean_para(para):
        for run in para.runs:
            if run.text:
                stats["cells"] += 1
                dirty = count_dirty_chars(run.text)
                if dirty:
                    stats["dirty"] += 1
                run.text = clean_text(run.text)

    for para in doc.paragraphs:
        clean_para(para)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    clean_para(para)

    doc.save(dst)
    return stats


PROCESSORS = {
    ".xlsx": process_xlsx,
    ".csv":  process_csv,
    ".txt":  process_text,
    ".md":   process_text,
    ".docx": process_docx,
}

# ── UI helpers ────────────────────────────────────────────────────────────────

BANNER = f"""
{C['cyan']}{C['bold']}╔══════════════════════════════════════════╗
║        TEXT CLEANER  —  Unicode Fixer    ║
║   Hỗ trợ: .xlsx  .csv  .txt  .md  .docx ║
╚══════════════════════════════════════════╝{C['reset']}
"""

def ask(prompt: str, default: str = "") -> str:
    try:
        ans = input(prompt).strip()
        return ans if ans else default
    except (KeyboardInterrupt, EOFError):
        print("\n" + c("yellow", "Đã huỷ."))
        sys.exit(0)


def resolve_output(src: Path) -> Path:
    print()
    print(c("bold", "Lưu kết quả:"))
    print(f"  {c('cyan', '1')}  Tạo file mới  {c('gray', f'({src.stem}_cleaned{src.suffix})')}")
    print(f"  {c('cyan', '2')}  Ghi đè file gốc")
    choice = ask(c("bold", "Chọn (1/2) [mặc định: 1]: "), "1")

    if choice == "2":
        confirm = ask(c("yellow", f"⚠  Ghi đè '{src.name}'? Không thể hoàn tác. (y/N): "), "n")
        if confirm.lower() != "y":
            print(c("gray", "→ Đổi sang tạo file mới."))
            choice = "1"

    if choice == "2":
        return src
    else:
        return src.parent / f"{src.stem}_cleaned{src.suffix}"


def print_stats(stats: dict, src: Path, dst: Path):
    label = "dòng" if src.suffix in (".txt", ".md") else "ô"
    print()
    print(c("green", "✔  Hoàn thành!"))
    print(f"   Đã quét  : {stats['cells']:,} {label}")
    print(f"   Có nhiễu : {c('yellow', str(stats['dirty'])) if stats['dirty'] else c('green', '0')} {label}")
    if dst != src:
        print(f"   File mới  : {c('cyan', str(dst))}")
    else:
        print(f"   Đã lưu   : {c('cyan', str(dst))}")


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    print(BANNER)

    # Nhận file từ argument hoặc hỏi interactive
    if len(sys.argv) > 1:
        file_path = Path(sys.argv[1])
    else:
        raw = ask(c("bold", "Nhập đường dẫn file cần xử lý: "))
        # Bỏ dấu nháy nếu user kéo thả file vào terminal
        raw = raw.strip("'\"")
        file_path = Path(raw)

    # Validate
    if not file_path.exists():
        print(c("red", f"\n✘  Không tìm thấy file: {file_path}"))
        sys.exit(1)

    ext = file_path.suffix.lower()
    if ext not in PROCESSORS:
        print(c("red", f"\n✘  Loại file '{ext}' chưa được hỗ trợ."))
        print(c("gray", f"   Hỗ trợ: {', '.join(PROCESSORS.keys())}"))
        sys.exit(1)

    print(f"\n   File    : {c('cyan', file_path.name)}")
    print(f"   Loại   : {c('gray', ext)}")
    print(f"   Kích thước: {c('gray', f'{file_path.stat().st_size / 1024:.1f} KB')}")

    dst = resolve_output(file_path)

    # Backup nếu ghi đè
    if dst == file_path:
        backup = file_path.with_suffix(file_path.suffix + ".bak")
        shutil.copy2(file_path, backup)
        print(c("gray", f"\n   Backup  : {backup.name}"))

    print(c("gray", "\n   Đang xử lý..."))

    try:
        stats = PROCESSORS[ext](file_path, dst)
    except ImportError as e:
        print(c("red", f"\n✘  {e}"))
        sys.exit(1)
    except Exception as e:
        print(c("red", f"\n✘  Lỗi khi xử lý file: {e}"))
        sys.exit(1)

    print_stats(stats, file_path, dst)
    print()


if __name__ == "__main__":
    main()
